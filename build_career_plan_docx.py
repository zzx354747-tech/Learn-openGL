from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT = "应用统计学个人职业生涯规划书.docx"


def set_cell_shading(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_width(cell, width):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_w = tc_pr.first_child_found_in("w:tcW")
    if tc_w is None:
        tc_w = OxmlElement("w:tcW")
        tc_pr.append(tc_w)
    tc_w.set(qn("w:w"), str(width))
    tc_w.set(qn("w:type"), "dxa")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in {"top": top, "start": start, "bottom": bottom, "end": end}.items():
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_run_font(run, name="宋体", size=None, bold=None, color=None):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)
    if size:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


def add_paragraph(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.first_line_indent = Pt(22)
    p.paragraph_format.line_spacing = 1.18
    p.paragraph_format.space_after = Pt(6)
    r = p.add_run(text)
    set_run_font(r, "宋体", 11)
    return p


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after = Pt(5)
    r = p.add_run(text)
    set_run_font(r, "黑体", 15 if level == 1 else 12, True, "1F4D78" if level == 1 else "2E74B5")
    return p


doc = Document()
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)
section.left_margin = Inches(1)
section.right_margin = Inches(1)

styles = doc.styles
normal = styles["Normal"]
normal.font.name = "宋体"
normal._element.rPr.rFonts.set(qn("w:eastAsia"), "宋体")
normal.font.size = Pt(11)

title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
title.paragraph_format.space_after = Pt(8)
run = title.add_run("个人职业生涯规划书")
set_run_font(run, "黑体", 18, True, "0B2545")

subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
subtitle.paragraph_format.space_after = Pt(14)
run = subtitle.add_run("专业：应用统计学")
set_run_font(run, "宋体", 11)

add_heading(doc, "一、引言")
add_paragraph(
    doc,
    "通过《大学生职业发展与就业指导（一）》的学习，我认识到职业生涯规划不是简单地选择一份工作，而是在自我认知、环境分析、目标设定和行动反馈之间持续循环的过程。作为应用统计学专业学生，我的专业学习既强调数学基础、统计推断和数据建模，也要求能够把数据方法应用到真实业务问题中。结合自己的兴趣、课程基础和就业环境，我希望把未来发展方向初步定位在数据分析与统计建模相关岗位，并通过大学阶段的持续训练提高专业能力、实践能力和职业适应力。"
)

add_heading(doc, "二、设定职业目标的过程")
add_paragraph(
    doc,
    "首先进行自我探索。我对概率论、数理统计、回归分析等课程兴趣较高，愿意通过数据发现问题、解释现象并提出建议；同时，我也意识到自己在编程熟练度、行业知识和表达展示方面仍有不足。根据霍兰德职业兴趣理论，我的特点更接近研究型与常规型，适合需要逻辑分析、数据处理和规范表达的工作。"
)
add_paragraph(
    doc,
    "其次分析外部环境。数字经济背景下，企事业单位对数据分析、市场研究、风险控制、运营分析和商业智能等岗位需求较多。应用统计学专业与这些岗位匹配度较高，但岗位竞争也要求学生掌握 Python、SQL、统计软件、可视化工具和一定业务理解能力。"
)
add_paragraph(
    doc,
    "最后运用生涯决策平衡单和加权评分法进行目标选择。我把备选方向设为数据分析师、统计建模/风控岗、考研深造三个选项，并从兴趣匹配、专业匹配、能力基础、就业机会、长期成长五个维度进行比较。权重越高说明该因素对我越重要，分值为 1 至 5 分。"
)

table = doc.add_table(rows=1, cols=7)
table.alignment = WD_TABLE_ALIGNMENT.CENTER
table.style = "Table Grid"
headers = ["职业选项", "兴趣\n25%", "专业\n25%", "能力\n20%", "机会\n15%", "成长\n15%", "加权结果"]
widths = [1800, 1100, 1100, 1100, 1100, 1100, 2060]
for i, h in enumerate(headers):
    cell = table.rows[0].cells[i]
    cell.text = h
    set_cell_width(cell, widths[i])
    set_cell_shading(cell, "F2F4F7")
    set_cell_margins(cell)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    for p in cell.paragraphs:
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for r in p.runs:
            set_run_font(r, "黑体", 9, True)

rows = [
    ["数据分析师", "5", "5", "4", "5", "4", "4.65"],
    ["统计建模/风控岗", "4", "5", "3", "4", "5", "4.20"],
    ["考研深造", "4", "5", "4", "3", "5", "4.25"],
]
for row in rows:
    cells = table.add_row().cells
    for i, value in enumerate(row):
        cells[i].text = value
        set_cell_width(cells[i], widths[i])
        set_cell_margins(cells[i])
        cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
        for p in cells[i].paragraphs:
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if i else WD_ALIGN_PARAGRAPH.LEFT
            for r in p.runs:
                set_run_font(r, "宋体", 9)

add_paragraph(
    doc,
    "根据评分结果，我的近期职业目标设定为：本科阶段以数据分析师为主要就业目标，同时保留统计建模和考研深造的可能性。具体来说，我希望在毕业前具备独立完成数据清洗、统计分析、可视化呈现和基础建模的能力，能够胜任互联网、金融、咨询或制造业中的初级数据分析岗位。"
)

add_heading(doc, "三、实现职业目标的具体行动")
add_paragraph(
    doc,
    "第一，夯实专业基础。认真学习概率论、数理统计、多元统计分析、时间序列、抽样调查等课程，形成较系统的统计思维；课后通过整理错题、复盘案例和阅读教材补充知识，避免只会套公式而不能解释结果。"
)
add_paragraph(
    doc,
    "第二，提高工具能力。以 Python 和 SQL 为重点，逐步掌握 pandas、numpy、matplotlib、sklearn 等常用工具，能够完成数据读取、清洗、描述分析、回归建模和结果可视化；同时熟悉 Excel、SPSS 或 R 语言，增强不同场景下的适应性。"
)
add_paragraph(
    doc,
    "第三，积累项目与实践。每学期至少完成一个小型数据项目，例如校园消费数据分析、城市房价影响因素分析、用户留存分析或问卷调查报告。项目过程要保留数据来源、分析思路、代码和结论，逐步形成个人作品集。"
)
add_paragraph(
    doc,
    "第四，提升通用能力。主动练习报告写作、PPT 展示和沟通表达，学会把模型结果转化为普通读者能理解的结论；同时关注行业案例，理解数据分析在运营、营销、风控和公共管理中的具体用途。"
)

add_heading(doc, "四、评析行动成果")
add_paragraph(
    doc,
    "从短期看，职业规划的行动成果可以通过课程成绩、工具熟练度、项目数量和实习反馈来评价。如果能够在专业课中保持稳定成绩，独立完成两到三个数据分析项目，并能清楚解释分析过程，说明阶段性行动是有效的。"
)
add_paragraph(
    doc,
    "同时，我也需要警惕只重视证书和软件操作、忽视统计原理与业务理解的问题。数据分析岗位真正需要的是提出问题、选择方法、解释结果和支持决策的综合能力。因此，评价成果不能只看学了多少工具，还要看能否把工具用于解决真实问题。"
)

add_heading(doc, "五、职业目标及行动的动态调整")
add_paragraph(
    doc,
    "职业目标不是一次确定后就不再变化。今后我会在每学期末进行一次复盘，重点检查兴趣是否稳定、能力短板是否改善、外部就业形势是否变化。如果发现自己对理论研究更有兴趣，并且数学基础和英语水平能够支撑继续深造，我会把考研作为主要目标；如果在实习中发现自己更适合金融风险管理或统计建模，则会增加机器学习、信用评分和金融基础知识的学习。"
)
add_paragraph(
    doc,
    "在行动上，我会采用“目标—行动—反馈—修正”的循环方式：目标保持方向感，行动保持可执行，反馈来自课程、项目、竞赛、实习和老师同学的评价，修正则体现在学习重点和时间分配上。这样既能避免盲目坚持不合适的目标，也能防止因短期困难轻易放弃。"
)

add_heading(doc, "六、结语")
add_paragraph(
    doc,
    "通过本次职业生涯规划，我更加明确了应用统计学专业与未来职业之间的联系，也认识到职业发展需要把自我认知、专业学习和社会需求结合起来。数据分析师是我当前较适合的阶段性目标，但实现这一目标需要长期积累和持续调整。接下来，我将以课程学习为基础，以工具训练和项目实践为抓手，在不断反馈中完善个人能力结构，努力把专业优势转化为就业竞争力和长期发展能力。"
)

doc.save(OUT)
print(OUT)
